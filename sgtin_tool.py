"""SGTIN AI 01-21 tool (экспериментальная версия).

Графический инструмент для пакетной обработки SGTIN кодов маркировки.

Возможности:
- УДАЛИТЬ AI 01-21 / Добавить AI 01-21 (строгая проверка 27/31).
- В 27 / В 31 — извлечение SGTIN из GS1-строки любого формата
  (с префиксами, доп. AI 91/92, и т.д.).
- ЭКРАН — замена экранированных последовательностей (\\u001d, <GS>, \\n и т.д.) на реальные символы.
- СОЗДАТЬ DATAMATRIX — генерация GS1 DataMatrix из кодов нижней панели.
- Выгрузка результата в Excel.
"""
from __future__ import annotations

import base64
import binascii
import html
import os
import re
import tempfile
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from typing import Iterable, List, Tuple

try:
    from openpyxl import Workbook
except ImportError:  # pragma: no cover - openpyxl is a runtime dep
    Workbook = None  # type: ignore[assignment]

try:
    import ppf.datamatrix as _ppfdm  # type: ignore[import]
except ImportError:  # pragma: no cover
    _ppfdm = None  # type: ignore[assignment]


# ---------------------------------------------------------------------------
# Поддержка FNC1 в ppf.datamatrix.
#
# Настоящий GS1 DataMatrix начинается со спецсимвола FNC1 (codeword 232 в
# ASCII-кодировке DataMatrix). Сама библиотека ppf.datamatrix его не умеет
# вставлять, поэтому мы переопределяем её ASCII-кодек: символ-маркер FNC1
# (\x1e, который в реальных кодах не встречается) кодируется в codeword 232.
# ---------------------------------------------------------------------------
FNC1_MARKER = "\x1e"  # внутренний маркер начала GS1-кода (превратится в 232)


def _register_fnc1_codec() -> None:
    if _ppfdm is None:
        return
    try:
        import codecs
        from ppf.datamatrix import codec_ascii as _ca
    except Exception:  # pragma: no cover
        return

    _digits = _ca.DIGITS

    def _encode_with_fnc1(msg):
        enc = []
        i = 0
        while i < len(msg):
            ch = msg[i]
            if ch == FNC1_MARKER:
                enc.append(232)  # codeword FNC1
                i += 1
                continue
            if ch in _digits and i + 1 < len(msg) and msg[i + 1] in _digits:
                enc.append(130 + int(msg[i:i + 2]))
                i += 2
                continue
            enc.append(list(ch.encode("ascii"))[0] + 1)
            i += 1
        return bytes(enc), len(enc)

    def _search(name):
        if name != "datamatrix.ascii":
            return None
        return codecs.CodecInfo(
            _encode_with_fnc1, _ca.decode_from_ascii, name="datamatrix.ascii"
        )

    codecs.register(_search)


_register_fnc1_codec()

try:
    from PIL import Image, ImageTk  # type: ignore[import]
except ImportError:  # pragma: no cover
    Image = None  # type: ignore[assignment]
    ImageTk = None  # type: ignore[assignment]


REMOVE_LEN = 31
ADD_LEN = 27
ADD_LEN_93 = 20  # GTIN(14) + Serial(6) — короткий формат для 93-кодов
GTIN_LEN = 14
SERIAL_LEN = 13
SERIAL_93_LEN = 6  # типичная длина Serial в кодах с AI(93)
AI_01 = "01"
AI_21 = "21"
AI_93 = "93"

# Префикс вида '[DATAMATRIX (GS1)]:' — мусор, который всегда вырезаем.
# Может встречаться несколько раз в одной строке.
DM_PREFIX_RE = re.compile(r"\[DATAMATRIX\s*\(GS1\)\]\s*:\s*", re.IGNORECASE)


def strip_dm_prefix(s: str) -> str:
    """Убрать все вхождения префикса '[DATAMATRIX (GS1)]:' из строки."""
    return DM_PREFIX_RE.sub(" ", s).strip()


def remove_ai(code: str) -> str:
    """Удалить AI(01) и AI(21) из SGTIN длиной 31 символ.

    Ожидается формат: "01" + GTIN(14) + "21" + Serial(13).
    Возвращает GTIN + Serial (27 символов).
    """
    if len(code) != REMOVE_LEN:
        raise ValueError(
            f"ожидалось {REMOVE_LEN} символов, получено {len(code)}"
        )
    if not code[: len(AI_01)] == AI_01:
        raise ValueError("строка не начинается с AI(01)")
    if code[len(AI_01) + GTIN_LEN : len(AI_01) + GTIN_LEN + len(AI_21)] != AI_21:
        raise ValueError("AI(21) не найден в ожидаемой позиции")
    gtin = code[len(AI_01) : len(AI_01) + GTIN_LEN]
    serial = code[len(AI_01) + GTIN_LEN + len(AI_21) :]
    return gtin + serial


def add_ai(code: str) -> str:
    """Добавить AI(01) и AI(21) к коду GTIN(14)+Serial.

    Принимает:
    - 27 символов: GTIN(14) + Serial(13)  → 01+GTIN+21+Serial (31)
    - 20 символов: GTIN(14) + Serial(6)   → 01+GTIN+21+Serial (24)  [93-формат]
    Возвращает "01" + GTIN + "21" + Serial.
    """
    if len(code) not in (ADD_LEN, ADD_LEN_93):
        raise ValueError(
            f"ожидалось {ADD_LEN} или {ADD_LEN_93} символов, получено {len(code)}"
        )
    gtin = code[:GTIN_LEN]
    serial = code[GTIN_LEN:]
    return AI_01 + gtin + AI_21 + serial


def cut_after_93(line: str) -> str:
    """Кнопка «93»: оставить только 01+GTIN(14)+21+Serial, вырезав всё с AI(93).

    Структура входа (разделитель GS перед 93 в «сыром» коде не пропечатан):
        01 + GTIN(14) + 21 + Serial(перем.) [GS|пробел] 93...<хвост>
    Serial бывает разной длины (например 6 или 13). Всё, что начинается
    с разделителя и AI(93), отбрасывается. Также снимается мусорный
    префикс '[DATAMATRIX (GS1)]:'.
    """
    s = strip_dm_prefix(line.strip()).strip()

    # Нормализуем явные разделители (GS/пробел/таб) — пригодится для поиска 93.
    # Сначала пытаемся отрезать по разделителю, за которым идёт 93.
    m = re.match(r"^(01\d{14}21.+?)[\x1d \t]+93\S*.*$", s)
    if m:
        return m.group(1)

    # Разделитель не пропечатан: 01+GTIN+21+Serial сразу слитно с 93...
    # Serial по структуре RU-маркировки — 6 или 13 символов. Пробуем оба.
    for serlen in (SERIAL_LEN, SERIAL_93_LEN):
        m = re.match(
            r"^(01\d{14}21\S{%d})93\S*$" % serlen, s
        )
        if m:
            return m.group(1)

    # Уже чистый код без 93 — проверяем, что это валидный 01+GTIN+21+Serial.
    m = re.match(r"^01\d{14}21\S+$", s)
    if m:
        return s

    raise ValueError("не удалось разобрать код (нет паттерна 01+GTIN+21+Serial[+93])")


def extract_sgtin(line: str) -> Tuple[str, str]:
    """Извлечь (GTIN14, Serial) из любого формата GS1-строки.

    Поддерживает:
    - 27 символов: GTIN(14) + Serial(13)
    - 31 символ: 01+GTIN(14)+21+Serial(13)
    - Длинные GS1-строки с доп. AI (91,92,...), разделённые
      пробелами или GS (\x1d).
    - Префиксы вроде '[DATAMATRIX (GS1)]:'
    """
    s = line.strip()

    # Убираем мусорный префикс '[DATAMATRIX (GS1)]:' (любое число вхождений).
    # ВАЖНО: двоеточие НЕ является разделителем внутри SGTIN — оно может быть
    # частью Serial (AI 21). Поэтому режем ТОЛЬКО полноценный префикс
    # '[DATAMATRIX (GS1)]:', а одиночные ':' внутри кода не трогаем.
    if DM_PREFIX_RE.search(s):
        s = strip_dm_prefix(s)
        # После снятия префикса могла остаться лишь первая запись с пробелами —
        # берём левую часть до первого пробела/GS для дальнейшего поиска.
    s = s.strip()

    # 27 символов и все печатные без пробелов
    if len(s) == ADD_LEN and " " not in s:
        return s[:GTIN_LEN], s[GTIN_LEN:]

    # 31 символ с AI
    if (
        len(s) == REMOVE_LEN
        and s.startswith(AI_01)
        and s[len(AI_01) + GTIN_LEN : len(AI_01) + GTIN_LEN + len(AI_21)] == AI_21
        and " " not in s
    ):
        return s[len(AI_01) : len(AI_01) + GTIN_LEN], s[len(AI_01) + GTIN_LEN + len(AI_21) :]

    # Общий поиск: 01<14цифр>21<13 непробельных>
    m = re.search(r"01(\d{14})21(\S{13})", s)
    if m:
        return m.group(1), m.group(2)

    # Попробуем поиск с переменной длиной serial (до разделителя)
    m = re.search(r"01(\d{14})21([^\x1d\s]+)", s)
    if m:
        serial = m.group(2)
        return m.group(1), serial[:SERIAL_LEN] if len(serial) > SERIAL_LEN else serial

    raise ValueError("не удалось извлечь SGTIN (нет паттерна 01+GTIN+21+Serial)")


def convert_to_27(line: str) -> str:
    """Извлечь SGTIN в 27-символьном формате (GTIN+Serial) из любого входа."""
    gtin, serial = extract_sgtin(line)
    return gtin + serial


def convert_to_31(line: str) -> str:
    """Извлечь SGTIN в 31-символьном формате (01+GTIN+21+Serial) из любого входа."""
    gtin, serial = extract_sgtin(line)
    return AI_01 + gtin + AI_21 + serial


_ESCAPE_RE = re.compile(
    r"\\(?:u([0-9a-fA-F]{4})|x([0-9a-fA-F]{2})|([nrtbfv'\"\\/ ]))"
)


def _escape_repl(m: re.Match) -> str:
    if m.group(1):  # \uXXXX
        return chr(int(m.group(1), 16))
    if m.group(2):  # \xHH
        return chr(int(m.group(2), 16))
    ch = m.group(3)
    return {
        "n": "\n", "r": "\r", "t": "\t", "b": "\b",
        "f": "\f", "v": "\v", "\\\\": "\\",
        "'": "'", '"': '"', "/": "/", " ": " ",
    }.get(ch, ch)


def unescape_line(line: str) -> str:
    """Заменить экранированные последовательности на реальные символы.

    Поддерживает:
    - HTML-сущности: &lt; &gt; &amp; &apos; &quot; (плюс остальные стандартные)
    - GS1-метки: <GS>, [GS], \\u001d → \\x1d
    - Python-эскейпы: \\n, \\t, \\r, \\xHH, \\uXXXX, \\\\
    """
    s = line
    # 1) HTML-сущности (то, о чём явно просил пользователь):
    #    &lt; &gt; &amp; &apos; &quot; и прочие стандартные.
    s = html.unescape(s)
    # 2) Общепринятые placeholder'ы для GS1-меток.
    s = s.replace("<GS>", "\x1d").replace("<gs>", "\x1d")
    s = s.replace("[GS]", "\x1d").replace("[gs]", "\x1d")
    # 3) Стандартные Python-escape последовательности.
    s = _ESCAPE_RE.sub(_escape_repl, s)
    return s


def insert_gs_separators(line: str, strict: bool = False) -> str:
    """Подготовить GS1-payload для DataMatrix: расставить разделители GS (\\x1d).

    Логика:
    1. Снимаем мусорный префикс '[DATAMATRIX (GS1)]:' (двоеточие внутри
       Serial при этом не трогается).
    2. Если в коде уже есть пробелы/табы/GS — каждый такой разделитель
       приводим к одному GS. Это «вариант по пробелу»: пользователь сам
       вставил полный код с пробелами после Serial и перед 91/92/93.
    3. Если разделителей нет — пытаемся разобрать «слитный» код по
       структуре RU-маркировки и вставить GS перед AI 91 / 92 / 93:
           01+GTIN(14)+21+Serial  [GS]  91+...  [GS]  92+...
           01+GTIN(14)+21+Serial  [GS]  93+...
    4. Короткий код (27/31, без доп. AI) — возвращаем как есть.

    strict=True: если в коде ЕСТЬ доп. AI (91/92/93), но структуру
    разобрать и корректно расставить GS не удалось — бросаем ValueError,
    чтобы не сгенерировать нечитаемый DataMatrix.
    """
    s = strip_dm_prefix(line.strip()).strip()
    if not s:
        return s

    has_extra_ai = bool(re.search(r"(?:^|[\x1d \t])9[123]", s)) or \
        bool(re.search(r"21.+9[123]", s))

    # 2) Если есть явные разделители (пробел/таб/GS) — нормализуем в GS.
    if re.search(r"[ \t\x1d]", s):
        s2 = re.sub(r"[ \t\x1d]+", "\x1d", s).strip("\x1d")
        # Проверяем, что начинается с корректного SGTIN.
        if strict and not re.match(r"^01\d{14}21", s2):
            raise ValueError("после нормализации код не начинается с 01+GTIN+21")
        return s2

    # 3) Разделителей нет — разбираем слитный код по структуре.
    # 01+GTIN+21+Serial(13) + 91+Key(4) + 92+Crypto(rest)
    m = re.match(r"^(01\d{14}21.{13})(91.{4})(92.+)$", s)
    if m:
        return m.group(1) + "\x1d" + m.group(2) + "\x1d" + m.group(3)
    # 01+GTIN+21+Serial(13) + 91+Key
    m = re.match(r"^(01\d{14}21.{13})(91.+)$", s)
    if m:
        return m.group(1) + "\x1d" + m.group(2)
    # 01+GTIN+21+Serial(13) + 92+Crypto
    m = re.match(r"^(01\d{14}21.{13})(92.+)$", s)
    if m:
        return m.group(1) + "\x1d" + m.group(2)
    # 01+GTIN+21+Serial(6 или 13) + 93+хвост
    for serlen in (SERIAL_93_LEN, SERIAL_LEN):
        m = re.match(r"^(01\d{14}21.{%d})(93.+)$" % serlen, s)
        if m:
            return m.group(1) + "\x1d" + m.group(2)

    # 4) Доп. AI присутствует, но структуру не распознали — в строгом
    # режиме это ошибка (иначе DataMatrix будет нечитаемым).
    if strict and has_extra_ai:
        raise ValueError(
            "есть AI 91/92/93, но не удалось расставить разделители GS"
        )
    return s


_B64_RE = re.compile(r"^[A-Za-z0-9+/]+={0,2}$")
# Признак уже готового кода маркировки: 01 + 14 цифр (GTIN).
_SGTIN_START_RE = re.compile(r"^\x1e?01\d{14}")


def looks_like_base64(s: str) -> bool:
    """Грубая эвристика: похожа ли строка на base64 (а не на код маркировки).

    Код маркировки начинается с '01'+14 цифр — такие строки base64 НЕ считаем.
    base64: только символы [A-Za-z0-9+/=], длина кратна 4, достаточно длинная.
    """
    s = s.strip()
    if not s or len(s) < 16 or len(s) % 4 != 0:
        return False
    # Если это уже готовый код маркировки — это не base64.
    if _SGTIN_START_RE.match(s):
        return False
    return bool(_B64_RE.match(s))


def try_decode_base64(line: str) -> Tuple[str, bool]:
    """Попробовать раскодировать base64-строку в код маркировки.

    Возвращает (код, was_base64). Если строка была base64 и декодировалась
    в осмысленный код маркировки (содержит 01+GTIN+21) — возвращаем его и
    True. Иначе возвращаем исходную строку и False.

    Внутри декодированного кода уже может присутствовать FNC1 (\\x1d или
    \\xe8) и разделители GS — мы их сохраняем как есть.
    """
    s = line.strip()
    if not looks_like_base64(s):
        return line, False
    try:
        raw = base64.b64decode(s, validate=True)
    except (binascii.Error, ValueError):
        return line, False
    # Пробуем расшифровать как текст (latin-1 покрывает все байты 0..255,
    # включая GS \x1d и FNC1 \xe8, не теряя данных).
    try:
        decoded = raw.decode("utf-8")
    except UnicodeDecodeError:
        decoded = raw.decode("latin-1")
    # Нормализуем возможный FNC1: байт 0xE8 (ASCII 232) или \x1d в начале —
    # это служебные символы GS1, при дальнейшей генерации учтём отдельно.
    # Проверяем, что внутри реально код маркировки.
    probe = decoded.lstrip("\x1d\xe8\x1e").strip()
    if re.match(r"^01\d{14}21", probe):
        return decoded, True
    return line, False


def has_embedded_fnc1(code: str) -> bool:
    """Есть ли в начале кода уже вставленный FNC1 (\\xe8 / \\x1e / маркер)."""
    return code[:1] in ("\xe8", "\x1e", FNC1_MARKER)


def matrix_to_image(matrix: List[List[int]], target_size: int = 300,
                    quiet: int = 1):
    """Отрисовать матрицу DataMatrix в чёрно-белую PIL.Image.

    Размер модуля подбирается под `target_size`, картинка центрируется
    на белом фоне `target_size × target_size`.
    """
    if Image is None:
        raise RuntimeError("Pillow не установлен (pip install Pillow)")
    rows = len(matrix)
    cols = len(matrix[0]) if matrix else 0
    total_r = rows + 2 * quiet
    total_c = cols + 2 * quiet
    module = max(1, target_size // max(total_r, total_c))
    w = module * total_c
    h = module * total_r
    img = Image.new("L", (w, h), 255)
    pixels = img.load()
    for r, row in enumerate(matrix):
        for c, v in enumerate(row):
            if v:
                xx = (c + quiet) * module
                yy = (r + quiet) * module
                for dy in range(module):
                    for dx in range(module):
                        pixels[xx + dx, yy + dy] = 0
    if (w, h) != (target_size, target_size):
        canvas = Image.new("L", (target_size, target_size), 255)
        canvas.paste(img, ((target_size - w) // 2, (target_size - h) // 2))
        return canvas
    return img


def make_datamatrix(payload: str, size: int = 300, fnc1: bool = True):
    """Сгенерировать DataMatrix-картинку для готового payload.

    fnc1=True — генерируется настоящий GS1 DataMatrix: в начало кода
    добавляется спецсимвол FNC1 (codeword 232). Это рекомендованный
    формат для Честного знака. fnc1=False — обычный DataMatrix без FNC1.
    """
    if _ppfdm is None:
        raise RuntimeError("ppf.datamatrix не установлен (pip install ppf.datamatrix)")
    data = (FNC1_MARKER + payload) if fnc1 else payload
    dm = _ppfdm.DataMatrix(data)
    return matrix_to_image(dm.matrix, target_size=size)


def process_lines(lines: Iterable[str], mode: str) -> Tuple[List[str], List[str]]:
    """Применить выбранную операцию к каждой непустой строке.

    Возвращает (результаты, ошибки). Каждая ошибка содержит номер строки
    (1-based) и описание проблемы.
    """
    ops = {
        "remove": remove_ai,
        "add": add_ai,
        "to_27": convert_to_27,
        "to_31": convert_to_31,
        "unescape": unescape_line,
        "cut_93": cut_after_93,
    }
    op = ops.get(mode)
    if op is None:
        raise ValueError(f"неизвестная операция: {mode}")

    results: List[str] = []
    errors: List[str] = []
    for idx, raw in enumerate(lines, start=1):
        code = raw.strip()
        if not code:
            continue
        try:
            results.append(op(code))
        except ValueError as exc:
            errors.append(f"строка {idx}: '{code}' — {exc}")
    return results, errors


# Цветовая палитра (Tailwind-подобная).
BG = "#f1f5f9"          # фон окна — slate-100
CARD = "#ffffff"        # фон текстовых полей
BORDER = "#cbd5e1"      # рамка / линии
TEXT = "#0f172a"        # основной текст
MUTED = "#475569"       # вторичный текст / статус-бар
DANGER = "#dc2626"      # «УДАЛИТЬ» — красный
DANGER_HOVER = "#b91c1c"
SUCCESS = "#16a34a"     # «Добавить» — зелёный
SUCCESS_HOVER = "#15803d"
NEUTRAL = "#64748b"     # «Очистить» — slate
NEUTRAL_HOVER = "#475569"
PRIMARY = "#2563eb"     # «Excel» — синий
PRIMARY_HOVER = "#1d4ed8"
AMBER = "#d97706"       # «ЭКРАН» — янтарный
AMBER_HOVER = "#b45309"
INDIGO = "#7c3aed"      # «СОЗДАТЬ DATAMATRIX» — фиолетовый
INDIGO_HOVER = "#6d28d9"
FONT_UI = ("Segoe UI", 10)
FONT_UI_BOLD = ("Segoe UI", 10, "bold")
FONT_HEADING = ("Segoe UI", 11, "bold")
FONT_BIG_BTN = ("Segoe UI", 12, "bold")
FONT_MONO = ("Consolas", 11)


class SgtinApp(tk.Tk):
    """Главное окно приложения."""

    def __init__(self) -> None:
        super().__init__()
        self.title("SGTIN AI 01-21 — инструмент")
        self.geometry("1000x760")
        self.minsize(700, 520)
        self.configure(bg=BG)

        # Буфер последних клавиатурных событий для диагностики.
        self._key_log: list[str] = []
        self._diag_text: tk.Text | None = None
        self._last_log_serial: int | None = None

        # Галочка GS1 DataMatrix (FNC1) — включена при каждом запуске.
        self.fnc1_var = tk.BooleanVar(value=True)

        self._setup_style()
        self._build_ui()
        self._build_menu()
        self._bind_shortcuts()

    # ------------------------------------------------------------------
    # Style
    # ------------------------------------------------------------------
    def _setup_style(self) -> None:
        style = ttk.Style(self)
        # «clam» позволяет красить ttk.Button фоном на всех платформах.
        if "clam" in style.theme_names():
            style.theme_use("clam")

        style.configure(".", background=BG, foreground=TEXT, font=FONT_UI)
        style.configure("TFrame", background=BG)
        style.configure("Card.TFrame", background=CARD, relief="flat")
        style.configure("TLabel", background=BG, foreground=TEXT, font=FONT_UI)
        style.configure(
            "Heading.TLabel", background=BG, foreground=TEXT, font=FONT_HEADING
        )
        style.configure(
            "Status.TLabel", background=BG, foreground=MUTED, font=FONT_UI
        )
        style.configure(
            "TCheckbutton", background=BG, foreground=TEXT, font=FONT_UI
        )
        style.map(
            "TCheckbutton",
            background=[("active", BG)],
        )

        def _btn(name: str, bg: str, hover: str) -> None:
            style.configure(
                name,
                background=bg,
                foreground="white",
                bordercolor=bg,
                lightcolor=bg,
                darkcolor=bg,
                focuscolor=bg,
                padding=(14, 10),
                font=FONT_UI_BOLD,
                borderwidth=0,
                relief="flat",
            )
            style.map(
                name,
                background=[("active", hover), ("pressed", hover)],
                foreground=[("disabled", "#e2e8f0")],
            )

        _btn("Danger.TButton", DANGER, DANGER_HOVER)
        _btn("Success.TButton", SUCCESS, SUCCESS_HOVER)
        _btn("Neutral.TButton", NEUTRAL, NEUTRAL_HOVER)
        _btn("Primary.TButton", PRIMARY, PRIMARY_HOVER)
        _btn("Amber.TButton", AMBER, AMBER_HOVER)

        # Крупная кнопка для «СОЗДАТЬ DATAMATRIX» — больше padding и более крупный шрифт.
        style.configure(
            "BigDM.TButton",
            background=INDIGO,
            foreground="white",
            bordercolor=INDIGO,
            lightcolor=INDIGO,
            darkcolor=INDIGO,
            focuscolor=INDIGO,
            padding=(18, 16),
            font=FONT_BIG_BTN,
            borderwidth=0,
            relief="flat",
        )
        style.map(
            "BigDM.TButton",
            background=[("active", INDIGO_HOVER), ("pressed", INDIGO_HOVER)],
            foreground=[("disabled", "#e2e8f0")],
        )

    # ------------------------------------------------------------------
    # UI
    # ------------------------------------------------------------------
    def _build_ui(self) -> None:
        root = ttk.Frame(self, padding=14, style="TFrame")
        root.pack(fill=tk.BOTH, expand=True)
        root.rowconfigure(1, weight=1)
        root.rowconfigure(4, weight=1)
        root.columnconfigure(0, weight=1)

        ttk.Label(
            root,
            text="Входные SGTIN (по одному в строке):",
            style="Heading.TLabel",
        ).grid(row=0, column=0, sticky="w", pady=(0, 6))
        self.input_text = self._make_text_area(root, row=1)

        middle = ttk.Frame(root, style="TFrame")
        middle.grid(row=2, column=0, sticky="ew", pady=14)
        for col in range(4):
            middle.columnconfigure(col, weight=1, uniform="btn")

        # Ряд 1 — преобразование
        ttk.Button(
            middle,
            text="УДАЛИТЬ AI 01-21",
            command=self.on_remove,
            style="Danger.TButton",
            takefocus=False,
        ).grid(row=0, column=0, padx=(0, 6), sticky="ew")
        ttk.Button(
            middle,
            text="Добавить AI 01-21 (93 тоже)",
            command=self.on_add,
            style="Success.TButton",
            takefocus=False,
        ).grid(row=0, column=1, padx=6, sticky="ew")
        ttk.Button(
            middle,
            text="В 27",
            command=self.on_to_27,
            style="Danger.TButton",
            takefocus=False,
        ).grid(row=0, column=2, padx=6, sticky="ew")
        ttk.Button(
            middle,
            text="В 31",
            command=self.on_to_31,
            style="Success.TButton",
            takefocus=False,
        ).grid(row=0, column=3, padx=(6, 0), sticky="ew")

        # Ряд 2 — утилиты + работа с AI(93)
        ttk.Button(
            middle,
            text="ЭКРАН",
            command=self.on_unescape,
            style="Amber.TButton",
            takefocus=False,
        ).grid(row=1, column=0, padx=(0, 6), pady=(8, 0), sticky="ew")
        ttk.Button(
            middle,
            text="93",
            command=self.on_cut_93,
            style="Danger.TButton",
            takefocus=False,
        ).grid(row=1, column=1, padx=6, pady=(8, 0), sticky="ew")
        ttk.Button(
            middle,
            text="Добавить 93",
            command=self.on_add,
            style="Success.TButton",
            takefocus=False,
        ).grid(row=1, column=2, padx=6, pady=(8, 0), sticky="ew")
        ttk.Button(
            middle,
            text="Очистить результат",
            command=self.on_clear_output,
            style="Neutral.TButton",
            takefocus=False,
        ).grid(row=1, column=3, padx=(6, 0), pady=(8, 0), sticky="ew")

        # Ряд 3 — большая выделенная кнопка генерации DataMatrix,
        # на всю ширину панели.
        ttk.Button(
            middle,
            text="СОЗДАТЬ DATAMATRIX",
            command=self.on_create_datamatrix,
            style="BigDM.TButton",
            takefocus=False,
        ).grid(row=2, column=0, columnspan=4, pady=(12, 0), sticky="ew")

        # Ряд 4 — галочка GS1 DataMatrix (FNC1).
        ttk.Checkbutton(
            middle,
            text="GS1 DataMatrix (добавлять FNC1) — рекомендуется для Честного знака",
            variable=self.fnc1_var,
            style="TCheckbutton",
            takefocus=False,
        ).grid(row=3, column=0, columnspan=4, pady=(8, 0), sticky="w")

        # Ряд 5 — выгрузка в Excel, в самом низу панели.
        ttk.Button(
            middle,
            text="Выгрузить в Excel",
            command=self.on_export_excel,
            style="Primary.TButton",
            takefocus=False,
        ).grid(row=4, column=0, columnspan=4, pady=(8, 0), sticky="ew")

        ttk.Label(root, text="Результат:", style="Heading.TLabel").grid(
            row=3, column=0, sticky="w", pady=(0, 6)
        )
        self.output_text = self._make_text_area(root, row=4)

        self.status_var = tk.StringVar(value="Готово")
        ttk.Label(
            root,
            textvariable=self.status_var,
            anchor="w",
            style="Status.TLabel",
        ).grid(row=5, column=0, sticky="ew", pady=(10, 0))

    def _make_text_area(self, parent: ttk.Frame, row: int) -> tk.Text:
        frame = ttk.Frame(parent, style="Card.TFrame")
        frame.grid(row=row, column=0, sticky="nsew")
        frame.rowconfigure(0, weight=1)
        frame.columnconfigure(0, weight=1)
        frame.configure(borderwidth=1, relief="solid")

        text = tk.Text(
            frame,
            wrap="none",
            undo=True,
            maxundo=-1,
            font=FONT_MONO,
            bg=CARD,
            fg=TEXT,
            insertbackground=TEXT,
            selectbackground=PRIMARY,
            selectforeground="white",
            relief="flat",
            borderwidth=0,
            padx=8,
            pady=6,
            highlightthickness=0,
        )
        text.grid(row=0, column=0, sticky="nsew")

        yscroll = ttk.Scrollbar(frame, orient="vertical", command=text.yview)
        yscroll.grid(row=0, column=1, sticky="ns")
        xscroll = ttk.Scrollbar(frame, orient="horizontal", command=text.xview)
        xscroll.grid(row=1, column=0, sticky="ew")
        text.configure(yscrollcommand=yscroll.set, xscrollcommand=xscroll.set)

        self._attach_context_menu(text)
        return text

    def _build_menu(self) -> None:
        menubar = tk.Menu(self)
        self.config(menu=menubar)

        edit = tk.Menu(menubar, tearoff=0)
        edit.add_command(
            label="Вырезать",
            command=lambda: self._invoke_on_focused("<<Cut>>"),
        )
        edit.add_command(
            label="Копировать",
            command=lambda: self._invoke_on_focused("<<Copy>>"),
        )
        edit.add_command(
            label="Вставить",
            command=lambda: self._invoke_on_focused("<<Paste>>"),
        )
        edit.add_separator()
        edit.add_command(
            label="Выделить всё",
            command=lambda: self._select_all_focused(),
        )
        edit.add_command(
            label="Удалить выделенное",
            command=lambda: self._delete_selection_focused(),
        )
        edit.add_separator()
        edit.add_command(
            label="Отменить",
            command=lambda: self._undo_focused(),
        )
        edit.add_command(
            label="Повторить",
            command=lambda: self._redo_focused(),
        )
        menubar.add_cascade(label="Правка", menu=edit)

        help_menu = tk.Menu(menubar, tearoff=0)
        help_menu.add_command(
            label="Диагностика клавиш",
            command=self._open_diagnostics,
        )
        menubar.add_cascade(label="Помощь", menu=help_menu)

    def _focused_text(self) -> tk.Text | None:
        w = self.focus_get()
        if isinstance(w, tk.Text):
            return w
        # Резерв — верхнее поле ввода.
        return self.input_text

    def _invoke_on_focused(self, virtual: str) -> None:
        w = self._focused_text()
        if w is not None:
            w.event_generate(virtual)

    def _select_all_focused(self) -> None:
        w = self._focused_text()
        if w is None:
            return
        w.tag_add("sel", "1.0", "end-1c")
        w.mark_set("insert", "1.0")
        w.see("insert")

    def _delete_selection_focused(self) -> None:
        w = self._focused_text()
        if w is None:
            return
        try:
            w.delete("sel.first", "sel.last")
        except tk.TclError:
            pass

    def _undo_focused(self) -> None:
        w = self._focused_text()
        if w is None:
            return
        try:
            w.edit_undo()
        except tk.TclError:
            pass

    def _redo_focused(self) -> None:
        w = self._focused_text()
        if w is None:
            return
        try:
            w.edit_redo()
        except tk.TclError:
            pass

    # ------------------------------------------------------------------
    # Диагностика клавиатуры
    # ------------------------------------------------------------------
    def _log_key_event(self, event: tk.Event) -> None:
        # Пишем в лог только если диагностика жива (иначе быстро разрастётся).
        if self._diag_text is None:
            return
        # Одно событие может прилететь и в <KeyPress>, и в <Control-KeyPress>;
        # дедуплицируем по event.serial.
        serial = getattr(event, "serial", None)
        if serial is not None and serial == self._last_log_serial:
            return
        self._last_log_serial = serial
        char = event.char
        char_repr = "''" if char == "" else repr(char)
        line = (
            f"keysym={event.keysym!r:>18}  keycode={event.keycode:>4}  "
            f"state=0x{event.state:08x}  char={char_repr}"
        )
        self._key_log.append(line)
        self._key_log = self._key_log[-200:]
        try:
            self._diag_text.insert("end", line + "\n")
            self._diag_text.see("end")
        except tk.TclError:
            self._diag_text = None

    def _open_diagnostics(self) -> None:
        if self._diag_text is not None:
            try:
                self._diag_text.winfo_toplevel().deiconify()
                self._diag_text.winfo_toplevel().lift()
                return
            except tk.TclError:
                self._diag_text = None

        win = tk.Toplevel(self)
        win.title("Диагностика клавиатуры")
        win.geometry("760x420")
        win.configure(bg=BG)

        ttk.Label(
            win,
            text=(
                "Сфокусируйтесь на верхнем/нижнем поле в главном окне и нажимайте Ctrl+А/С/М/Я — сюдат будет писаться, какой код приходит от Windows."
            ),
            style="Status.TLabel",
            wraplength=720,
        ).pack(fill="x", padx=10, pady=(10, 6))

        frame = ttk.Frame(win, style="Card.TFrame")
        frame.pack(fill="both", expand=True, padx=10, pady=(0, 10))
        frame.configure(borderwidth=1, relief="solid")
        frame.rowconfigure(0, weight=1)
        frame.columnconfigure(0, weight=1)

        log = tk.Text(
            frame,
            wrap="none",
            font=FONT_MONO,
            bg=CARD,
            fg=TEXT,
            relief="flat",
            borderwidth=0,
            padx=8,
            pady=6,
            highlightthickness=0,
        )
        log.grid(row=0, column=0, sticky="nsew")
        sb = ttk.Scrollbar(frame, orient="vertical", command=log.yview)
        sb.grid(row=0, column=1, sticky="ns")
        log.configure(yscrollcommand=sb.set)

        for line in self._key_log:
            log.insert("end", line + "\n")
        log.see("end")

        btns = ttk.Frame(win, style="TFrame")
        btns.pack(fill="x", padx=10, pady=(0, 10))
        ttk.Button(
            btns,
            text="Очистить",
            command=lambda: (self._key_log.clear(), log.delete("1.0", "end")),
            style="Neutral.TButton",
        ).pack(side="left")
        ttk.Button(
            btns,
            text="Скопировать в буфер",
            command=lambda: (
                self.clipboard_clear(),
                self.clipboard_append("\n".join(self._key_log)),
            ),
            style="Primary.TButton",
        ).pack(side="left", padx=(8, 0))

        def on_close() -> None:
            self._diag_text = None
            win.destroy()

        win.protocol("WM_DELETE_WINDOW", on_close)
        self._diag_text = log

    def _attach_context_menu(self, widget: tk.Text) -> None:
        menu = tk.Menu(widget, tearoff=0)
        menu.add_command(
            label="Вырезать", command=lambda: widget.event_generate("<<Cut>>")
        )
        menu.add_command(
            label="Копировать", command=lambda: widget.event_generate("<<Copy>>")
        )
        menu.add_command(
            label="Вставить", command=lambda: widget.event_generate("<<Paste>>")
        )
        menu.add_separator()
        menu.add_command(
            label="Выделить всё",
            command=lambda: widget.event_generate("<<SelectAll>>"),
        )
        menu.add_command(
            label="Удалить выделенное",
            command=lambda: self._delete_selection(widget),
        )

        def show(event: tk.Event) -> None:
            try:
                menu.tk_popup(event.x_root, event.y_root)
            finally:
                menu.grab_release()

        widget.bind("<Button-3>", show)

    @staticmethod
    def _delete_selection(widget: tk.Text) -> None:
        try:
            widget.delete("sel.first", "sel.last")
        except tk.TclError:
            pass

    # Windows virtual-key коды — НЕ зависят от раскладки.
    # На X11 hardware keycodes другие (38=A, 54=C, 55=V, 53=X, 29=Y, 52=Z
    # на типичной PC-клавиатуре), но тоже от раскладки не зависят.
    _CTRL_ACTIONS = {
        # Windows VK
        65: "select_all", 67: "copy", 86: "paste",
        88: "cut", 89: "redo", 90: "undo",
        # X11 hardware keycodes (PC AT-101)
        38: "select_all", 54: "copy", 55: "paste",
        53: "cut", 29: "redo", 52: "undo",
    }

    # Резервный матчинг по keysym — на случай если keycode не помог
    # (например, нестандартная клавиатура или X11 на другом железе).
    _CTRL_KEYSYMS = {
        # Латиница
        "a": "select_all", "A": "select_all",
        "c": "copy", "C": "copy",
        "v": "paste", "V": "paste",
        "x": "cut", "X": "cut",
        "y": "redo", "Y": "redo",
        "z": "undo", "Z": "undo",
        # Кириллица (для X11)
        "Cyrillic_ef": "select_all", "Cyrillic_EF": "select_all",
        "Cyrillic_es": "copy", "Cyrillic_ES": "copy",
        "Cyrillic_em": "paste", "Cyrillic_EM": "paste",
        "Cyrillic_che": "cut", "Cyrillic_CHE": "cut",
        "Cyrillic_en": "redo", "Cyrillic_EN": "redo",
        "Cyrillic_ya": "undo", "Cyrillic_YA": "undo",
    }

    def _bind_shortcuts(self) -> None:
        # На Windows Tkinter при русской раскладке в комбинациях с Ctrl event.keysym
        # часто приходит как "??", и явные биндинги вроде <Control-Cyrillic_es>
        # не срабатывают. Даже <Control-KeyPress> может не вызваться
        # из-за разбора секвенции. Поэтому подвязываемся на сырой <KeyPress>
        # и сами смотрим на event.state для Ctrl. Используем event.keycode
        # (на Windows это VK-код — от раскладки не зависит).
        for widget in (self.input_text, self.output_text):
            widget.bind("<KeyPress>", self._on_any_keypress, add="+")
            widget.bind("<Control-KeyPress>", self._on_ctrl_key)
            widget.bind("<Control-Insert>", self._copy)
            widget.bind("<Shift-Insert>", self._paste)
            widget.bind("<Shift-Delete>", self._cut)
            widget.bind("<Delete>", self._on_delete)
            widget.bind("<BackSpace>", self._on_backspace)

    def _on_any_keypress(self, event: tk.Event) -> str | None:
        # Резервный путь: срабатывает на любой клавише; сами проверяем
        # бит Ctrl в state. Нужно потому, что Tk на Windows при RU-раскладке
        # может не вызывать <Control-KeyPress>, если keysym не распознан.
        # Когда открыто окно диагностики — логируем всё, чтобы видеть,
        # что вообще приходит от ОС.
        if self._diag_text is not None:
            self._log_key_event(event)
        if not (event.state & 0x0004):  # Ctrl не зажат
            return None
        return self._on_ctrl_key(event)

    def _on_ctrl_key(self, event: tk.Event) -> str | None:
        # Логируем в диагностику (если открыта) — всегда, даже если действие не найдено.
        self._log_key_event(event)

        # ВАЖНО про event.state на Windows: его биты могут включать
        # «постоянно установленные» флаги вроде Caps/NumLock и побочные
        # бит-маски, которые **не** соответствуют классическим X11
        # ShiftMask/ControlMask/Mod1Mask. Поэтому НЕ полагаемся на состояние
        # как фильтр — диспетчеризуем по самому надёжному сигналу:
        # event.char для Ctrl+латиница == control-символ \x01..\x1a (Tk
        # формирует его и под русской раскладкой, потому что физическая
        # клавиша одна и та же), event.keycode (Windows VK — независим от
        # раскладки) и event.keysym (Latin/Cyrillic_*) как резерв.
        action: str | None = None

        char = event.char or ""
        if len(char) == 1 and 1 <= ord(char) <= 26:
            # Ctrl+<буква>: \x01==A, \x03==C, \x16==V, \x18==X, \x19==Y,
            # \x1a==Z. Маппим в латинский keysym.
            letter = chr(ord(char) + ord("a") - 1)
            action = self._CTRL_KEYSYMS.get(letter)

        if action is None:
            action = self._CTRL_ACTIONS.get(event.keycode)
        if action is None:
            action = self._CTRL_KEYSYMS.get(event.keysym)

        if action is None:
            return None  # не наш шорткат — дать сработать стандартному поведению

        shift = bool(event.state & 0x0001)
        if action == "undo" and shift:
            action = "redo"  # Ctrl+Shift+Z = redo

        if action == "select_all":
            return self._select_all(event)
        if action == "copy":
            return self._copy(event)
        if action == "paste":
            return self._paste(event)
        if action == "cut":
            return self._cut(event)
        if action == "undo":
            return self._undo(event)
        if action == "redo":
            return self._redo(event)
        return None

    @staticmethod
    def _copy(event: tk.Event) -> str:
        event.widget.event_generate("<<Copy>>")
        return "break"

    @staticmethod
    def _paste(event: tk.Event) -> str:
        event.widget.event_generate("<<Paste>>")
        return "break"

    @staticmethod
    def _cut(event: tk.Event) -> str:
        event.widget.event_generate("<<Cut>>")
        return "break"

    @staticmethod
    def _select_all(event: tk.Event) -> str:
        widget: tk.Text = event.widget  # type: ignore[assignment]
        widget.tag_add("sel", "1.0", "end-1c")
        widget.mark_set("insert", "1.0")
        widget.see("insert")
        return "break"

    @staticmethod
    def _undo(event: tk.Event) -> str:
        widget: tk.Text = event.widget  # type: ignore[assignment]
        try:
            widget.edit_undo()
        except tk.TclError:
            pass
        return "break"

    @staticmethod
    def _redo(event: tk.Event) -> str:
        widget: tk.Text = event.widget  # type: ignore[assignment]
        try:
            widget.edit_redo()
        except tk.TclError:
            pass
        return "break"

    @staticmethod
    def _on_delete(event: tk.Event) -> str | None:
        widget: tk.Text = event.widget  # type: ignore[assignment]
        try:
            widget.delete("sel.first", "sel.last")
            return "break"
        except tk.TclError:
            return None  # нет выделения — пусть отработает стандартный Delete

    @staticmethod
    def _on_backspace(event: tk.Event) -> str | None:
        widget: tk.Text = event.widget  # type: ignore[assignment]
        try:
            widget.delete("sel.first", "sel.last")
            return "break"
        except tk.TclError:
            return None

    # ------------------------------------------------------------------
    # Actions
    # ------------------------------------------------------------------
    def _get_input_lines(self) -> List[str]:
        raw = self.input_text.get("1.0", "end-1c")
        return raw.splitlines()

    def _set_output(self, results: List[str]) -> None:
        self.output_text.delete("1.0", "end")
        if results:
            self.output_text.insert("1.0", "\n".join(results))

    def _run(self, mode: str) -> None:
        lines = self._get_input_lines()
        if not any(line.strip() for line in lines):
            messagebox.showwarning(
                "Нет данных", "Введите хотя бы один SGTIN в верхнем поле."
            )
            return

        results, errors = process_lines(lines, mode)
        self._set_output(results)

        if errors:
            preview = "\n".join(errors[:20])
            extra = "" if len(errors) <= 20 else f"\n… и ещё {len(errors) - 20}"
            messagebox.showerror(
                "Ошибки в данных",
                f"Не обработано строк: {len(errors)}\n\n{preview}{extra}",
            )
        self.status_var.set(
            f"Обработано: {len(results)}; ошибок: {len(errors)}"
        )

    def on_remove(self) -> None:
        self._run("remove")

    def on_add(self) -> None:
        self._run("add")

    def on_to_27(self) -> None:
        self._run("to_27")

    def on_to_31(self) -> None:
        self._run("to_31")

    def on_unescape(self) -> None:
        self._run("unescape")

    def on_cut_93(self) -> None:
        self._run("cut_93")

    def on_clear_output(self) -> None:
        self.output_text.delete("1.0", "end")
        self.status_var.set("Результат очищен")

    def on_export_excel(self) -> None:
        if Workbook is None:
            messagebox.showerror(
                "openpyxl не установлен",
                "Установите пакет openpyxl: pip install openpyxl",
            )
            return

        data = self.output_text.get("1.0", "end-1c").splitlines()
        data = [line for line in data if line.strip()]
        if not data:
            messagebox.showwarning(
                "Нет данных", "Нижнее поле пустое — нечего выгружать."
            )
            return

        path = filedialog.asksaveasfilename(
            title="Сохранить как",
            defaultextension=".xlsx",
            filetypes=[("Excel Workbook", "*.xlsx")],
            initialfile="sgtin_export.xlsx",
        )
        if not path:
            return

        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "SGTIN"
            ws.append(["SGTIN"])
            for row in data:
                # Сохраняем как текст, чтобы Excel не превращал в число
                # и не терял ведущие нули.
                ws.append([str(row)])
            for cell in ws["A"]:
                cell.number_format = "@"
            wb.save(path)
        except OSError as exc:
            messagebox.showerror("Ошибка записи", str(exc))
            return

        self.status_var.set(
            f"Выгружено {len(data)} строк в {os.path.basename(path)}"
        )
        messagebox.showinfo(
            "Готово", f"Сохранено строк: {len(data)}\nФайл: {path}"
        )

    # ------------------------------------------------------------------
    # DataMatrix
    # ------------------------------------------------------------------
    def on_create_datamatrix(self) -> None:
        if _ppfdm is None or Image is None:
            messagebox.showerror(
                "Зависимости не установлены",
                "Для генерации DataMatrix нужны пакеты "
                "ppf.datamatrix и Pillow.\n"
                "Установите: pip install ppf.datamatrix Pillow",
            )
            return

        data = self.output_text.get("1.0", "end-1c").splitlines()
        codes = [line.strip() for line in data if line.strip()]
        if not codes:
            messagebox.showwarning(
                "Нет данных",
                "Нижнее поле пустое — нечего кодировать.",
            )
            return

        # Готовим payload каждого кода.
        # 1) Если строка похожа на base64 — декодируем её в код маркировки.
        # 2) Вставляем GS-разделители (strict — ошибка, если 91/92/93 без структуры).
        # 3) Если в декодированном коде УЖЕ есть FNC1 — для него добавление
        #    FNC1 отключаем индивидуально (чтобы не задвоить).
        global_fnc1 = self.fnc1_var.get()
        prepared: List[str] = []
        fnc1_flags: List[bool] = []
        good_codes: List[str] = []
        errors: List[str] = []
        for idx, c in enumerate(codes, start=1):
            try:
                decoded, was_b64 = try_decode_base64(c)
                embedded = has_embedded_fnc1(decoded)
                # Убираем уже вставленный FNC1-байт из начала перед обработкой
                # (мы добавим свой при генерации, если нужно).
                clean = decoded
                if embedded:
                    clean = decoded.lstrip("\xe8\x1e" + FNC1_MARKER)
                payload = insert_gs_separators(clean, strict=True)
                prepared.append(payload)
                # FNC1 ставим, если включена галочка. Если код пришёл из
                # base64 с уже встроенным FNC1 — всё равно ставим один свой
                # (после очистки), сохраняя единообразие. Если галочка выкл —
                # не ставим.
                fnc1_flags.append(global_fnc1)
                good_codes.append(c)
            except ValueError as exc:
                errors.append(f"строка {idx}: '{c}' — {exc}")

        if errors:
            preview = "\n".join(errors[:20])
            extra = "" if len(errors) <= 20 else f"\n… и ещё {len(errors) - 20}"
            messagebox.showerror(
                "Некорректные коды для DataMatrix",
                f"Пропущено строк: {len(errors)}\n\n{preview}{extra}",
            )

        if not good_codes:
            self.status_var.set("DataMatrix: нет корректных кодов")
            return

        codes = good_codes

        if len(codes) == 1:
            DataMatrixPreview(
                self, payload=prepared[0], display=codes[0],
                fnc1=fnc1_flags[0],
            )
            self.status_var.set("Окно DataMatrix открыто")
            return

        self._export_datamatrix_batch(codes, prepared, fnc1_flags=fnc1_flags)

    def _export_datamatrix_batch(
        self, codes: List[str], prepared: List[str],
        fnc1_flags: List[bool] | None = None,
    ) -> None:
        # Спрашиваем формат — Excel или Word.
        choice = _ask_choice(
            self,
            title="Выгрузка DataMatrix",
            text=(
                f"Найдено {len(codes)} кодов.\n"
                "Куда выгрузить картинки DataMatrix?"
            ),
            options=[("Excel (.xlsx)", "excel"), ("Word (.docx)", "word")],
        )
        if not choice:
            return

        if choice == "excel":
            ext = ".xlsx"
            ftypes = [("Excel Workbook", "*.xlsx")]
            default = "datamatrix_export.xlsx"
        else:
            ext = ".docx"
            ftypes = [("Word Document", "*.docx")]
            default = "datamatrix_export.docx"

        path = filedialog.asksaveasfilename(
            title="Сохранить как",
            defaultextension=ext,
            filetypes=ftypes,
            initialfile=default,
        )
        if not path:
            return

        try:
            # Генерируем картинки один раз.
            if fnc1_flags is None:
                fnc1_flags = [True] * len(prepared)
            images = [
                make_datamatrix(p, size=300, fnc1=f)
                for p, f in zip(prepared, fnc1_flags)
            ]
        except Exception as exc:  # pragma: no cover
            messagebox.showerror(
                "Ошибка генерации DataMatrix", str(exc)
            )
            return

        try:
            if choice == "excel":
                _export_xlsx(path, codes, images)
            else:
                _export_docx(path, codes, images)
        except Exception as exc:  # pragma: no cover
            messagebox.showerror("Ошибка записи", str(exc))
            return

        self.status_var.set(
            f"Сохранено {len(codes)} DataMatrix в {os.path.basename(path)}"
        )
        messagebox.showinfo(
            "Готово",
            f"Сохранено кодов: {len(codes)}\nФайл: {path}",
        )


# ----------------------------------------------------------------------
# DataMatrix helpers (preview window, export, clipboard, print)
# ----------------------------------------------------------------------
def _ask_choice(parent: tk.Misc, title: str, text: str,
                options: List[Tuple[str, str]]) -> str | None:
    """Показать модальный диалог с кнопками-вариантами; вернуть выбранное value."""
    win = tk.Toplevel(parent)
    win.title(title)
    win.configure(bg=BG)
    win.transient(parent)
    win.grab_set()
    win.resizable(False, False)

    frame = ttk.Frame(win, padding=18, style="TFrame")
    frame.pack(fill=tk.BOTH, expand=True)
    ttk.Label(frame, text=text, style="TLabel", justify="left").pack(
        anchor="w", pady=(0, 12)
    )

    btn_frame = ttk.Frame(frame, style="TFrame")
    btn_frame.pack(fill=tk.X)

    result: dict[str, str | None] = {"value": None}

    def _pick(value: str) -> None:
        result["value"] = value
        win.destroy()

    for label, value in options:
        ttk.Button(
            btn_frame,
            text=label,
            style="Primary.TButton",
            command=lambda v=value: _pick(v),
            takefocus=False,
        ).pack(side=tk.LEFT, padx=(0, 8))

    ttk.Button(
        btn_frame,
        text="Отмена",
        style="Neutral.TButton",
        command=win.destroy,
        takefocus=False,
    ).pack(side=tk.RIGHT)

    # Центрируем окно над родителем.
    win.update_idletasks()
    px = parent.winfo_rootx() + parent.winfo_width() // 2 - win.winfo_width() // 2
    py = parent.winfo_rooty() + parent.winfo_height() // 2 - win.winfo_height() // 2
    win.geometry(f"+{max(0, px)}+{max(0, py)}")

    parent.wait_window(win)
    return result["value"]


def _copy_image_to_clipboard_windows(img) -> None:
    """Положить картинку в системный буфер обмена Windows (через ctypes).

    На не-Windows бросает RuntimeError.
    """
    import ctypes
    import io

    if os.name != "nt":
        raise RuntimeError(
            "Копирование картинки в буфер обмена доступно только в Windows."
        )

    # BMP-формат для CF_DIB — пропускаем 14-байтный BITMAPFILEHEADER.
    buf = io.BytesIO()
    img.convert("RGB").save(buf, format="BMP")
    data = buf.getvalue()[14:]
    buf.close()

    user32 = ctypes.windll.user32
    kernel32 = ctypes.windll.kernel32
    CF_DIB = 8
    GMEM_MOVEABLE = 0x0002

    if not user32.OpenClipboard(0):
        raise RuntimeError("OpenClipboard failed")
    try:
        user32.EmptyClipboard()
        h_mem = kernel32.GlobalAlloc(GMEM_MOVEABLE, len(data))
        if not h_mem:
            raise RuntimeError("GlobalAlloc failed")
        ptr = kernel32.GlobalLock(h_mem)
        try:
            ctypes.memmove(ptr, data, len(data))
        finally:
            kernel32.GlobalUnlock(h_mem)
        if not user32.SetClipboardData(CF_DIB, h_mem):
            raise RuntimeError("SetClipboardData failed")
    finally:
        user32.CloseClipboard()


def _print_image_windows(img) -> None:
    """Открыть системный диалог печати для картинки (Windows)."""
    if os.name != "nt":
        raise RuntimeError("Печать доступна только в Windows.")
    fd, tmp_path = tempfile.mkstemp(prefix="datamatrix_", suffix=".png")
    os.close(fd)
    img.save(tmp_path)
    os.startfile(tmp_path, "print")  # type: ignore[attr-defined]


def _export_xlsx(path: str, codes: List[str], images: List["Image.Image"],
                 img_px: int = 300) -> None:
    from openpyxl import Workbook
    from openpyxl.drawing.image import Image as XLImage

    wb = Workbook()
    ws = wb.active
    ws.title = "DataMatrix"
    ws.append(["SGTIN", "DataMatrix"])
    ws.column_dimensions["A"].width = 50
    # Excel column width в «символах»: ширина изображения ~ width * 7px.
    ws.column_dimensions["B"].width = img_px / 7 + 2

    tmp_files: List[str] = []
    try:
        for i, (code, img) in enumerate(zip(codes, images), start=2):
            ws.cell(row=i, column=1, value=str(code))
            ws.cell(row=i, column=1).number_format = "@"

            fd, tmp = tempfile.mkstemp(prefix="dm_", suffix=".png")
            os.close(fd)
            img.save(tmp)
            tmp_files.append(tmp)

            xl_img = XLImage(tmp)
            xl_img.width = img_px
            xl_img.height = img_px
            ws.add_image(xl_img, f"B{i}")
            # Высота строки в «пунктах»: ~ height * 0.75.
            ws.row_dimensions[i].height = img_px * 0.75
        wb.save(path)
    finally:
        for f in tmp_files:
            try:
                os.unlink(f)
            except OSError:
                pass


def _export_docx(path: str, codes: List[str], images: List["Image.Image"],
                 img_px: int = 300) -> None:
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    doc.add_heading("SGTIN DataMatrix", level=1)

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "SGTIN"
    hdr[1].text = "DataMatrix"

    tmp_files: List[str] = []
    try:
        inch = img_px / 96  # px → дюймы при 96 dpi
        for code, img in zip(codes, images):
            row = table.add_row().cells
            row[0].text = str(code)

            fd, tmp = tempfile.mkstemp(prefix="dm_", suffix=".png")
            os.close(fd)
            img.save(tmp)
            tmp_files.append(tmp)

            run = row[1].paragraphs[0].add_run()
            run.add_picture(tmp, width=Inches(inch))
        doc.save(path)
    finally:
        for f in tmp_files:
            try:
                os.unlink(f)
            except OSError:
                pass


class DataMatrixPreview(tk.Toplevel):
    """Окно предварительного просмотра одиночного DataMatrix.

    Кнопки: Сохранить PNG, Копировать в буфер обмена, Печать, ×2 (увеличить).
    """

    def __init__(self, parent: tk.Misc, payload: str, display: str,
                 fnc1: bool = True) -> None:
        super().__init__(parent)
        self.title("DataMatrix")
        self.configure(bg=BG)
        self.resizable(False, False)
        self.transient(parent if isinstance(parent, tk.Wm) else None)

        self.payload = payload
        self.display = display
        self.fnc1 = fnc1
        self.size = 300
        self._photo: ImageTk.PhotoImage | None = None
        self._image = None  # PIL.Image
        self._build()
        self._render()

    def _build(self) -> None:
        frame = ttk.Frame(self, padding=14, style="TFrame")
        frame.pack(fill=tk.BOTH, expand=True)

        ttk.Label(
            frame,
            text=self.display,
            style="Status.TLabel",
            wraplength=520,
            justify="left",
        ).pack(anchor="w", pady=(0, 10))

        self._img_label = tk.Label(frame, bg=CARD, bd=1, relief="solid")
        self._img_label.pack()

        btns = ttk.Frame(frame, style="TFrame")
        btns.pack(fill=tk.X, pady=(14, 0))

        ttk.Button(
            btns, text="Сохранить PNG", style="Primary.TButton",
            command=self._save_png, takefocus=False,
        ).pack(side=tk.LEFT, padx=(0, 6))
        ttk.Button(
            btns, text="Копировать", style="Neutral.TButton",
            command=self._copy_to_clipboard, takefocus=False,
        ).pack(side=tk.LEFT, padx=6)
        ttk.Button(
            btns, text="Печать", style="Neutral.TButton",
            command=self._print, takefocus=False,
        ).pack(side=tk.LEFT, padx=6)
        self._zoom_btn = ttk.Button(
            btns, text="Увеличить ×2", style="Amber.TButton",
            command=self._toggle_zoom, takefocus=False,
        )
        self._zoom_btn.pack(side=tk.RIGHT)

    def _render(self) -> None:
        self._image = make_datamatrix(self.payload, size=self.size, fnc1=self.fnc1)
        self._photo = ImageTk.PhotoImage(self._image)  # type: ignore[union-attr]
        self._img_label.configure(
            image=self._photo, width=self.size, height=self.size
        )

    def _toggle_zoom(self) -> None:
        self.size = 600 if self.size == 300 else 300
        self._zoom_btn.configure(
            text="Уменьшить ×1" if self.size == 600 else "Увеличить ×2"
        )
        self._render()

    def _save_png(self) -> None:
        path = filedialog.asksaveasfilename(
            title="Сохранить как",
            defaultextension=".png",
            filetypes=[("PNG image", "*.png")],
            initialfile="datamatrix.png",
            parent=self,
        )
        if not path:
            return
        try:
            self._image.save(path)
        except OSError as exc:
            messagebox.showerror("Ошибка записи", str(exc), parent=self)
            return
        messagebox.showinfo("Готово", f"Сохранено: {path}", parent=self)

    def _copy_to_clipboard(self) -> None:
        try:
            _copy_image_to_clipboard_windows(self._image)
        except RuntimeError as exc:
            messagebox.showinfo("Копирование", str(exc), parent=self)
        except Exception as exc:  # pragma: no cover
            messagebox.showerror("Не удалось скопировать", str(exc), parent=self)

    def _print(self) -> None:
        try:
            _print_image_windows(self._image)
        except RuntimeError as exc:
            messagebox.showinfo("Печать", str(exc), parent=self)
        except Exception as exc:  # pragma: no cover
            messagebox.showerror("Ошибка печати", str(exc), parent=self)


def main() -> None:
    app = SgtinApp()
    app.mainloop()


if __name__ == "__main__":
    main()
